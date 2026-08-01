# -*- coding: utf-8 -*-
"""
Nöroloji Klinik Asistan Botu
=============================
- Servis Vizit Notu: hasta bilgisi + ses kaydı -> yapılandırılmış SOAP notu
- Acil Konsültasyon: hasta bilgisi + ses kaydı -> yapılandırılmış konsültasyon notu
- Notlar SQLite'ta saklanır, /bugun ile günlük özet ve .docx export alınabilir
- /ara <isim> ile hasta bazlı arama yapılabilir

Kurulum ve deploy talimatları için README.md dosyasına bakın.
"""
import os
import logging
from datetime import date, datetime
from io import BytesIO

from dotenv import load_dotenv
load_dotenv()  # Railway'de ortam değişkenleri panelden gelir; yerelde .env kullanılır

from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup, InputFile
from telegram.ext import (
    Application, CommandHandler, MessageHandler, CallbackQueryHandler,
    ConversationHandler, ContextTypes, filters,
)
from openai import OpenAI
from docx import Document

import db
from prompts import VIZIT_SYSTEM_PROMPT, KONSULTASYON_SYSTEM_PROMPT

logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO,
)
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Ayarlar
# ---------------------------------------------------------------------------
TELEGRAM_BOT_TOKEN = os.environ["TELEGRAM_BOT_TOKEN"]
GROQ_API_KEY = os.environ["GROQ_API_KEY"]
# Botu yalnızca bu Telegram kullanıcı ID'sine izin verecek şekilde kısıtlar.
# Birden fazla kullanıcıya izin vermek için virgülle ayırıp liste yapabilirsiniz.
ALLOWED_USER_IDS = {
    int(uid) for uid in os.environ.get("ALLOWED_USER_ID", "").split(",") if uid.strip()
}

# Hem transkripsiyon (ses -> yazı) hem de not yapılandırma (yazı -> klinik not)
# Groq'un OpenAI-uyumlu, ücretsiz katmanı üzerinden yapılır. Kredi kartı gerekmez.
groq_client = OpenAI(api_key=GROQ_API_KEY, base_url="https://api.groq.com/openai/v1")

# Konuşma (conversation) durumları
(MENU, VIZIT_PATIENT, VIZIT_VOICE,
 KONSULT_PATIENT, KONSULT_VOICE, SEARCH_QUERY) = range(6)

# Tıbbi terminolojide doğruluk önceliği olduğundan "turbo" değil standart
# large-v3 modeli kullanılıyor.
TRANSCRIBE_MODEL = "whisper-large-v3"
# Groq'un ücretsiz katmanındaki güçlü, genel amaçlı modeli.
CHAT_MODEL = "llama-3.3-70b-versatile"


# ---------------------------------------------------------------------------
# Yetkilendirme
# ---------------------------------------------------------------------------
async def _check_auth(update: Update) -> bool:
    user_id = update.effective_user.id
    if ALLOWED_USER_IDS and user_id not in ALLOWED_USER_IDS:
        await update.effective_message.reply_text(
            "⛔ Bu bot yalnızca yetkili kullanıcılar için yapılandırılmıştır. "
            f"Telegram kullanıcı ID'niz: {user_id}"
        )
        logger.warning("Yetkisiz erişim denemesi: %s", user_id)
        return False
    return True


# ---------------------------------------------------------------------------
# Menü
# ---------------------------------------------------------------------------
def main_menu_keyboard():
    keyboard = [
        [InlineKeyboardButton("🏥 Servis Vizit Notu", callback_data="vizit")],
        [InlineKeyboardButton("🚨 Acil Konsültasyon", callback_data="konsultasyon")],
        [InlineKeyboardButton("📋 Bugünün Notları", callback_data="bugun")],
        [InlineKeyboardButton("🔍 Hasta Ara", callback_data="ara")],
    ]
    return InlineKeyboardMarkup(keyboard)


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    if not await _check_auth(update):
        return ConversationHandler.END
    await update.message.reply_text(
        "👋 Nöroloji Klinik Asistan Botuna hoş geldiniz.\n\n"
        "Ne yapmak istersiniz?",
        reply_markup=main_menu_keyboard(),
    )
    return MENU


async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await update.effective_message.reply_text(
        "İşlem iptal edildi. Menüye dönmek için /start yazabilirsiniz."
    )
    context.user_data.clear()
    return ConversationHandler.END


async def menu_router(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    query = update.callback_query
    await query.answer()
    choice = query.data

    if choice == "vizit":
        await query.edit_message_text(
            "🏥 *Servis Vizit Notu*\n\n"
            "Lütfen hasta adını veya oda numarasını yazın "
            "(örn: `Oda 304 - Ahmet Y.`):",
            parse_mode="Markdown",
        )
        return VIZIT_PATIENT

    if choice == "konsultasyon":
        await query.edit_message_text(
            "🚨 *Acil Konsültasyon*\n\n"
            "Lütfen hasta bilgisini yazın "
            "(örn: `Acil - Ayşe K., 67y, K`):",
            parse_mode="Markdown",
        )
        return KONSULT_PATIENT

    if choice == "bugun":
        await _send_todays_notes(update, context)
        return MENU

    if choice == "ara":
        await query.edit_message_text("🔍 Aramak istediğiniz hasta adını yazın:")
        return SEARCH_QUERY

    return MENU


# ---------------------------------------------------------------------------
# Servis Vizit Notu akışı
# ---------------------------------------------------------------------------
async def vizit_patient_received(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data["patient_info"] = update.message.text.strip()
    context.user_data["note_type"] = "vizit"
    await update.message.reply_text(
        f"✅ Hasta: {context.user_data['patient_info']}\n\n"
        "🎙️ Şimdi vizit muayenesini/notunu sesli olarak dikte edip gönderin."
    )
    return VIZIT_VOICE


async def vizit_voice_received(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await _process_voice_note(update, context, VIZIT_SYSTEM_PROMPT, "vizit")
    return MENU


# ---------------------------------------------------------------------------
# Acil Konsültasyon akışı
# ---------------------------------------------------------------------------
async def konsult_patient_received(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data["patient_info"] = update.message.text.strip()
    context.user_data["note_type"] = "konsultasyon"
    await update.message.reply_text(
        f"✅ Hasta: {context.user_data['patient_info']}\n\n"
        "🎙️ Şimdi konsültasyon öyküsünü ve nörolojik muayeneyi sesli olarak "
        "dikte edip gönderin."
    )
    return KONSULT_VOICE


async def konsult_voice_received(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    await _process_voice_note(update, context, KONSULTASYON_SYSTEM_PROMPT, "konsultasyon")
    return MENU


# ---------------------------------------------------------------------------
# Ortak: ses işleme -> transkripsiyon -> yapılandırma -> kayıt -> yanıt
# ---------------------------------------------------------------------------
async def _process_voice_note(update: Update, context: ContextTypes.DEFAULT_TYPE,
                               system_prompt: str, note_type: str):
    if not update.message.voice and not update.message.audio:
        await update.message.reply_text(
            "⚠️ Lütfen bir sesli mesaj (voice) veya ses dosyası gönderin."
        )
        return

    processing_msg = await update.message.reply_text("⏳ Ses kaydı işleniyor...")

    try:
        # 1) Ses dosyasını indir
        tg_file = await (update.message.voice or update.message.audio).get_file()
        audio_bytes = BytesIO()
        await tg_file.download_to_memory(out=audio_bytes)
        audio_bytes.seek(0)
        audio_bytes.name = "recording.ogg"  # Whisper API uzantıdan format algılar

        # 2) Transkripsiyon (Whisper - Groq üzerinden, hızlı ve ucuz)
        await processing_msg.edit_text("⏳ Ses metne çevriliyor...")
        transcript_resp = groq_client.audio.transcriptions.create(
            model=TRANSCRIBE_MODEL,
            file=audio_bytes,
            language="tr",
        )
        transcript = transcript_resp.text.strip()

        # 3) Yapılandırma (GPT ile klinik nota dönüştürme)
        await processing_msg.edit_text("⏳ Klinik not oluşturuluyor...")
        patient_info = context.user_data.get("patient_info", "[belirtilmedi]")
        now_str = datetime.now().strftime("%d.%m.%Y %H:%M")

        completion = groq_client.chat.completions.create(
            model=CHAT_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": (
                    f"Hasta bilgisi: {patient_info}\n"
                    f"Tarih/saat: {now_str}\n\n"
                    f"Ham transkript:\n{transcript}"
                )},
            ],
            temperature=0.2,
        )
        structured_note = completion.choices[0].message.content.strip()

        # 4) Kaydet
        note_id = db.save_note(
            note_type=note_type,
            patient_info=patient_info,
            transcript=transcript,
            structured_note=structured_note,
            telegram_user_id=update.effective_user.id,
        )

        # 5) Sonucu gönder
        await processing_msg.delete()
        await update.message.reply_text(
            f"✅ *Not #{note_id} kaydedildi*\n\n{structured_note}",
            parse_mode="Markdown",
        )
        await update.message.reply_text(
            "Başka bir işlem yapmak için /start yazabilirsiniz.",
            reply_markup=main_menu_keyboard(),
        )

    except Exception as e:
        logger.exception("Ses işleme hatası")
        await processing_msg.edit_text(
            f"❌ Bir hata oluştu: {e}\n\nLütfen tekrar deneyin veya /start ile yeniden başlayın."
        )
    finally:
        context.user_data.clear()


# ---------------------------------------------------------------------------
# Bugünün notları + arama
# ---------------------------------------------------------------------------
async def _send_todays_notes(update: Update, context: ContextTypes.DEFAULT_TYPE):
    notes = db.get_notes_by_date(date.today())
    target_message = update.callback_query.message if update.callback_query else update.message

    if not notes:
        await target_message.reply_text("Bugün için kayıtlı not bulunamadı.")
        return

    lines = [f"📋 *{date.today().strftime('%d.%m.%Y')} - Bugünün Notları* ({len(notes)} kayıt)\n"]
    for note_id, created_at, note_type, patient_info, _ in notes:
        saat = created_at.split("T")[1][:5]
        tip = "🏥 Vizit" if note_type == "vizit" else "🚨 Konsültasyon"
        lines.append(f"#{note_id} - {saat} - {tip} - {patient_info}")
    lines.append("\nDetayını görmek için: /not <id>\nWord olarak indirmek için: /export")

    await target_message.reply_text("\n".join(lines), parse_mode="Markdown")


async def get_note_detail(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not context.args:
        await update.message.reply_text("Kullanım: /not <id>  (örn: /not 12)")
        return
    try:
        note_id = int(context.args[0])
    except ValueError:
        await update.message.reply_text("Geçersiz ID.")
        return

    row = db.get_note_by_id(note_id)
    if not row:
        await update.message.reply_text("Bu ID ile kayıtlı not bulunamadı.")
        return
    _, created_at, note_type, patient_info, _, structured_note = row
    await update.message.reply_text(structured_note, parse_mode="Markdown")


async def export_today_docx(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not await _check_auth(update):
        return
    notes = db.get_notes_by_date(date.today())
    if not notes:
        await update.message.reply_text("Bugün için kayıtlı not bulunamadı.")
        return

    doc = Document()
    doc.add_heading(f"Nöroloji Notları - {date.today().strftime('%d.%m.%Y')}", level=1)
    for note_id, created_at, note_type, patient_info, structured_note in notes:
        saat = created_at.split("T")[1][:5]
        doc.add_heading(f"#{note_id} - {saat} - {patient_info}", level=2)
        for paragraph in structured_note.split("\n"):
            doc.add_paragraph(paragraph)
        doc.add_page_break()

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"notlar_{date.today().isoformat()}.docx"
    await update.message.reply_document(InputFile(buf, filename=filename))


async def search_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not await _check_auth(update):
        return
    await update.message.reply_text("🔍 Aramak istediğiniz hasta adını yazın:")
    return SEARCH_QUERY


async def search_query_received(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    query_text = update.message.text.strip()
    results = db.search_notes_by_patient(query_text)
    if not results:
        await update.message.reply_text("Sonuç bulunamadı.")
        return MENU

    lines = [f"🔍 *'{query_text}' için {len(results)} sonuç:*\n"]
    for note_id, created_at, note_type, patient_info, _ in results:
        tarih = created_at.split("T")[0]
        tip = "🏥" if note_type == "vizit" else "🚨"
        lines.append(f"#{note_id} - {tarih} - {tip} - {patient_info}")
    lines.append("\nDetay için: /not <id>")

    await update.message.reply_text("\n".join(lines), parse_mode="Markdown")
    return MENU


# ---------------------------------------------------------------------------
# Uygulama kurulumu
# ---------------------------------------------------------------------------
def main():
    db.init_db()
    app = Application.builder().token(TELEGRAM_BOT_TOKEN).build()

    conv_handler = ConversationHandler(
        entry_points=[CommandHandler("start", start)],
        states={
            MENU: [CallbackQueryHandler(menu_router)],
            VIZIT_PATIENT: [MessageHandler(filters.TEXT & ~filters.COMMAND, vizit_patient_received)],
            VIZIT_VOICE: [MessageHandler(filters.VOICE | filters.AUDIO, vizit_voice_received)],
            KONSULT_PATIENT: [MessageHandler(filters.TEXT & ~filters.COMMAND, konsult_patient_received)],
            KONSULT_VOICE: [MessageHandler(filters.VOICE | filters.AUDIO, konsult_voice_received)],
            SEARCH_QUERY: [MessageHandler(filters.TEXT & ~filters.COMMAND, search_query_received)],
        },
        fallbacks=[CommandHandler("cancel", cancel), CommandHandler("start", start)],
    )

    app.add_handler(conv_handler)
    app.add_handler(CommandHandler("not", get_note_detail))
    app.add_handler(CommandHandler("export", export_today_docx))
    app.add_handler(CommandHandler("bugun", _send_todays_notes))

    logger.info("Bot başlatılıyor (polling)...")
    app.run_polling(allowed_updates=Update.ALL_TYPES)


if __name__ == "__main__":
    main()

